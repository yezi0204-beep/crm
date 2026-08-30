"""附件解析模块。

支持从采购公告附件中提取文本内容：
- PDF：使用 pypdf
- Word (.docx)：使用 python-docx
- Excel (.xlsx)：使用 openpyxl
- 其他格式：记录文件名，不解析
"""
import os
import tempfile
import logging

logger = logging.getLogger(__name__)

# 文件扩展名 → 解析器映射
_PARSERS = {}

_MAX_FILE_SIZE = 20 * 1024 * 1024  # 20MB


def register_parser(extensions, parser_fn):
    """注册文件解析器。"""
    if isinstance(extensions, str):
        extensions = [extensions]
    for ext in extensions:
        _PARSERS[ext.lower()] = parser_fn


def parse_attachment(file_path: str) -> str:
    """解析附件文件，返回提取的文本内容。

    根据文件扩展名选择对应解析器。
    超过 20MB 的文件跳过解析。
    """
    if not file_path or not os.path.exists(file_path):
        return ''

    file_size = os.path.getsize(file_path)
    if file_size > _MAX_FILE_SIZE:
        logger.warning(f'附件过大({file_size}字节)，跳过: {file_path}')
        return ''

    ext = os.path.splitext(file_path)[1].lower()
    parser_fn = _PARSERS.get(ext)
    if not parser_fn:
        logger.info(f'不支持的附件格式: {ext}')
        return ''

    try:
        return parser_fn(file_path)
    except Exception as e:
        logger.error(f'解析附件失败 {file_path}: {e}')
        return ''


def _parse_pdf(file_path: str) -> str:
    """使用 pypdf 提取 PDF 文本。"""
    try:
        from pypdf import PdfReader
    except ImportError:
        logger.warning('pypdf 未安装，无法解析 PDF')
        return ''

    reader = PdfReader(file_path)
    texts = []
    for page in reader.pages:
        text = page.extract_text() or ''
        texts.append(text)
    return '\n'.join(texts).strip()


def _parse_docx(file_path: str) -> str:
    """使用 python-docx 提取 Word 文本。"""
    try:
        import docx
    except ImportError:
        logger.warning('python-docx 未安装，无法解析 Word')
        return ''

    doc = docx.Document(file_path)
    texts = []
    for para in doc.paragraphs:
        if para.text.strip():
            texts.append(para.text)

    # 提取表格内容
    for table in doc.tables:
        for row in table.rows:
            row_text = ' | '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                texts.append(row_text)

    return '\n'.join(texts).strip()


def _parse_xlsx(file_path: str) -> str:
    """使用 openpyxl 提取 Excel 文本。"""
    try:
        import openpyxl
    except ImportError:
        logger.warning('openpyxl 未安装，无法解析 Excel')
        return ''

    wb = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
    texts = []
    for ws in wb.worksheets:
        for row in ws.iter_rows(values_only=True):
            row_text = ' | '.join(str(cell).strip() for cell in row if cell)
            if row_text.strip():
                texts.append(row_text)
    wb.close()
    return '\n'.join(texts).strip()


# 注册解析器
register_parser('.pdf', _parse_pdf)
register_parser('.docx', _parse_docx)
register_parser('.xlsx', _parse_xlsx)
register_parser('.xls', _parse_xlsx)  # openpyxl 对 .xls 支持有限，降级处理


def download_and_parse(url: str, source_name: str = '', save_dir: str = None) -> dict:
    """下载附件并解析文本。

    返回 {path, text, filename} 或空 dict。
    """
    import requests

    try:
        from routes.leads import _get_search_session
        session = _get_search_session()
    except ImportError:
        session = requests.Session()

    try:
        resp = session.get(url, timeout=30, headers={
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) Chrome/120.0.0.0',
        })
        if resp.status_code != 200 or not resp.content:
            return {}

        # 从 URL 或响应头提取文件名
        filename = _extract_filename(url, resp)
        ext = os.path.splitext(filename)[1].lower()
        if ext not in _PARSERS:
            return {}

        # 保存到临时文件
        if not save_dir:
            save_dir = tempfile.gettempdir()
        os.makedirs(save_dir, exist_ok=True)
        file_path = os.path.join(save_dir, filename)
        with open(file_path, 'wb') as f:
            f.write(resp.content)

        text = parse_attachment(file_path)
        return {'path': file_path, 'text': text, 'filename': filename}
    except Exception as e:
        logger.error(f'下载附件失败 {url}: {e}')
        return {}


def _extract_filename(url: str, resp) -> str:
    """从 URL 或响应头提取文件名。"""
    # 优先从 Content-Disposition 提取
    cd = resp.headers.get('Content-Disposition', '')
    if cd:
        match = __import__('re').search(r'filename[*]?=["\']?([^"\';\s]+)', cd)
        if match:
            return os.path.basename(match.group(1))

    # 从 URL 提取
    from urllib.parse import urlparse, unquote
    parsed = urlparse(url)
    filename = unquote(os.path.basename(parsed.path))
    if filename and '.' in filename:
        return filename

    return 'attachment.bin'
