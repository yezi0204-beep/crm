"""知识库扩展路由 - 文档管理、批量导入、文件解析、CRM数据同步。

核心功能：
1. 文档库CRUD - 管理各类知识文档（拜访纪要/合同/投标文件/技术方案）
2. 批量导入 - 支持文件上传和批量文本导入
3. 文件解析 - 支持PDF/Word/Excel/TXT等格式
4. CRM数据同步 - 自动导入CRM系统数据关联
5. 语义搜索 - 基于向量索引的智能检索
6. 资质管理 - 人员资质和企业资质CRUD
7. AI智能分析 - 自动提取关键发现、客户需求、下一步行动、风险提示
"""
import os
import json
import uuid
import tempfile
from datetime import datetime

from flask import request, jsonify, current_app
from werkzeug.utils import secure_filename

from extensions import get_db, token_required, record_operation_log
from vector_search import (
    semantic_search, hybrid_search, index_document,
    generate_embedding, rebuild_all_vectors
)
from ai_analyzer import analyze_document, batch_analyze

from . import knowledge_ext_bp


UPLOAD_DIR = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), 'uploads', 'knowledge')
os.makedirs(UPLOAD_DIR, exist_ok=True)

ALLOWED_EXTENSIONS = {'txt', 'pdf', 'doc', 'docx', 'xls', 'xlsx', 'md', 'csv'}

DOC_TYPE_MAP = {
    'visit_summary': '拜访纪要',
    'contract': '合同',
    'bid_document': '投标文件',
    'technical_plan': '技术方案',
    'personnel_qualification': '人员资质',
    'company_qualification': '企业资质',
    'customer_info': '客户资料',
    'industry_report': '行业报告',
    'meeting_minutes': '会议纪要',
    'other': '其他'
}


def _allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


def _parse_file(file_path, mime_type):
    """解析文件内容，返回文本。

    针对 docx：很多企业文档（如拜访纪要）正文全部在表格中，
    仅取 paragraphs 会丢失正文，因此同时提取段落和表格内容。
    """
    ext = os.path.splitext(file_path)[1].lower()

    try:
        if ext == '.txt' or ext == '.md' or ext == '.csv':
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read()

        elif ext == '.pdf':
            try:
                from PyPDF2 import PdfReader
                reader = PdfReader(file_path)
                text = ''
                for page in reader.pages:
                    text += page.extract_text() or ''
                return text
            except ImportError:
                return None

        elif ext in ('.doc', '.docx'):
            try:
                from docx import Document
                doc = Document(file_path)
                parts = []
                # 段落
                for para in doc.paragraphs:
                    t = para.text.strip()
                    if t:
                        parts.append(t)
                # 表格（很多文档正文在表格中，如拜访纪要）
                for table in doc.tables:
                    for row in table.rows:
                        row_texts = []
                        for cell in row.cells:
                            ct = cell.text.strip()
                            # 合并单元格会在多个 cell 中重复，按行去重
                            if ct and ct not in row_texts:
                                row_texts.append(ct)
                        if row_texts:
                            parts.append(' | '.join(row_texts))
                return '\n'.join(parts)
            except ImportError:
                return None

        elif ext in ('.xls', '.xlsx'):
            try:
                import openpyxl
                wb = openpyxl.load_workbook(file_path, read_only=True)
                text = ''
                for sheet in wb.sheetnames:
                    ws = wb[sheet]
                    for row in ws.iter_rows(values_only=True):
                        text += ' '.join([str(cell) for cell in row if cell]) + '\n'
                return text
            except ImportError:
                return None

    except Exception:
        return None

    return None


@knowledge_ext_bp.route('/api/knowledge/documents', methods=['GET'])
@token_required
def get_documents():
    """获取知识库文档列表。"""
    data = request.current_user
    page = request.args.get('page', 1, type=int)
    per_page = request.args.get('per_page', 20, type=int)
    doc_type = request.args.get('doc_type', '')
    cust_id = request.args.get('cust_id', type=int)
    business_id = request.args.get('business_id', type=int)
    keyword = request.args.get('keyword', '')
    processed = request.args.get('processed', type=int)

    db = get_db()
    cursor = db.cursor()

    conditions = []
    params = []

    if doc_type:
        conditions.append("doc_type = ?")
        params.append(doc_type)
    if cust_id:
        conditions.append("cust_id = ?")
        params.append(cust_id)
    if business_id:
        conditions.append("business_id = ?")
        params.append(business_id)
    if keyword:
        conditions.append("(title LIKE ? OR content LIKE ? OR tags LIKE ?)")
        kw = f'%{keyword}%'
        params.extend([kw, kw, kw])
    if processed is not None:
        conditions.append("processed = ?")
        params.append(processed)

    where_clause = "WHERE " + " AND ".join(conditions) if conditions else ""

    cursor.execute(f"SELECT COUNT(*) as total FROM knowledge_documents {where_clause}", params)
    total = cursor.fetchone()['total']

    offset = (page - 1) * per_page
    cursor.execute(f"""
        SELECT d.*, c.company as customer_company, b.title as business_title,
               u.name as owner_name
        FROM knowledge_documents d
        LEFT JOIN customers c ON d.cust_id = c.id
        LEFT JOIN business b ON d.business_id = b.id
        LEFT JOIN users u ON d.owner_id = u.username
        {where_clause}
        ORDER BY d.created_at DESC
        LIMIT ? OFFSET ?
    """, params + [per_page, offset])

    docs = [dict(row) for row in cursor.fetchall()]

    for doc in docs:
        doc['doc_type_display'] = DOC_TYPE_MAP.get(doc.get('doc_type', ''), doc.get('doc_type', ''))
        # 解析轻量分析信息
        if doc.get('analysis_result'):
            try:
                a = json.loads(doc['analysis_result'])
                doc['analysis_summary'] = a.get('summary', '')[:100] if a.get('summary') else ''
                doc['analysis_tags'] = ','.join(a.get('tags', [])[:3])
            except:
                doc['analysis_summary'] = ''
                doc['analysis_tags'] = ''
        else:
            doc['analysis_summary'] = ''
            doc['analysis_tags'] = ''

    return jsonify({
        'code': 200,
        'message': 'success',
        'data': {
            'items': docs,
            'total': total,
            'page': page,
            'per_page': per_page,
            'doc_types': [{'value': k, 'label': v} for k, v in DOC_TYPE_MAP.items()]
        }
    })


@knowledge_ext_bp.route('/api/knowledge/documents', methods=['POST'])
@token_required
def create_document():
    """创建知识库文档。"""
    data = request.current_user
    username = data.get('username', '')
    req_data = request.get_json(silent=True) or {}

    doc_type = req_data.get('doc_type', 'other')
    title = req_data.get('title', '').strip()
    content = req_data.get('content', '').strip()
    cust_id = req_data.get('cust_id')
    business_id = req_data.get('business_id')
    contract_id = req_data.get('contract_id')
    tags = req_data.get('tags', '')
    summary = req_data.get('summary', '')
    doc_metadata = req_data.get('metadata', '')

    if not title:
        return jsonify({'code': 400, 'message': '标题不能为空', 'data': None})

    db = get_db()
    cursor = db.cursor()

    cursor.execute("""
        INSERT INTO knowledge_documents
        (doc_type, title, content, cust_id, business_id, contract_id,
         owner_id, tags, summary, doc_metadata, processed)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, 1)
    """, (doc_type, title, content, cust_id, business_id, contract_id,
          username, tags, summary, json.dumps(doc_metadata, ensure_ascii=False) if isinstance(doc_metadata, dict) else doc_metadata))
    db.commit()

    doc_id = cursor.lastrowid

    try:
        if content:
            index_document(doc_id, f"{title}\n{content}")
    except Exception:
        pass

    record_operation_log(username, '创建', '知识库文档', f'创建文档：{title}')

    return jsonify({
        'code': 200,
        'message': '创建成功',
        'data': {'id': doc_id, 'title': title}
    })


@knowledge_ext_bp.route('/api/knowledge/documents/<int:doc_id>', methods=['GET'])
@token_required
def get_document(doc_id):
    """获取单个文档详情。"""
    db = get_db()
    cursor = db.cursor()

    cursor.execute("""
        SELECT d.*, c.company as customer_company, b.title as business_title,
               u.name as owner_name
        FROM knowledge_documents d
        LEFT JOIN customers c ON d.cust_id = c.id
        LEFT JOIN business b ON d.business_id = b.id
        LEFT JOIN users u ON d.owner_id = u.username
        WHERE d.id = ?
    """, (doc_id,))
    doc = cursor.fetchone()
    if not doc:
        return jsonify({'code': 404, 'message': '文档不存在', 'data': None})

    result = dict(doc)
    result['doc_type_display'] = DOC_TYPE_MAP.get(result.get('doc_type', ''), result.get('doc_type', ''))

    # 解析分析结果
    if result.get('analysis_result'):
        try:
            result['analysis'] = json.loads(result['analysis_result'])
        except:
            result['analysis'] = None
    else:
        result['analysis'] = None

    return jsonify({
        'code': 200,
        'message': 'success',
        'data': result
    })


@knowledge_ext_bp.route('/api/knowledge/documents/<int:doc_id>', methods=['PUT'])
@token_required
def update_document(doc_id):
    """更新知识库文档。"""
    data = request.current_user
    username = data.get('username', '')
    req_data = request.get_json(silent=True) or {}

    db = get_db()
    cursor = db.cursor()

    cursor.execute("SELECT id FROM knowledge_documents WHERE id = ?", (doc_id,))
    if not cursor.fetchone():
        return jsonify({'code': 404, 'message': '文档不存在', 'data': None})

    fields = []
    params = []

    for field in ('title', 'content', 'doc_type', 'cust_id', 'business_id',
                  'contract_id', 'tags', 'summary', 'doc_metadata'):
        if field in req_data:
            fields.append(f"{field} = ?")
            val = req_data[field]
            if field == 'doc_metadata' and isinstance(val, dict):
                val = json.dumps(val, ensure_ascii=False)
            params.append(val)

    if not fields:
        return jsonify({'code': 400, 'message': '无更新字段', 'data': None})

    fields.append("updated_at = CURRENT_TIMESTAMP")
    params.append(doc_id)

    cursor.execute(f"UPDATE knowledge_documents SET {', '.join(fields)} WHERE id = ?", params)
    db.commit()

    if 'content' in req_data or 'title' in req_data:
        try:
            cursor.execute("SELECT title, content FROM knowledge_documents WHERE id = ?", (doc_id,))
            row = cursor.fetchone()
            index_document(doc_id, f"{row['title']}\n{row['content']}")
        except Exception:
            pass

    record_operation_log(username, '更新', '知识库文档', f'更新文档ID：{doc_id}')

    return jsonify({
        'code': 200,
        'message': '更新成功',
        'data': {'id': doc_id}
    })


@knowledge_ext_bp.route('/api/knowledge/documents/<int:doc_id>', methods=['DELETE'])
@token_required
def delete_document(doc_id):
    """删除知识库文档。"""
    data = request.current_user
    username = data.get('username', '')

    db = get_db()
    cursor = db.cursor()

    cursor.execute("SELECT id, file_path FROM knowledge_documents WHERE id = ?", (doc_id,))
    doc = cursor.fetchone()
    if not doc:
        return jsonify({'code': 404, 'message': '文档不存在', 'data': None})

    if doc['file_path'] and os.path.exists(doc['file_path']):
        try:
            os.remove(doc['file_path'])
        except Exception:
            pass

    cursor.execute("DELETE FROM knowledge_documents WHERE id = ?", (doc_id,))
    cursor.execute("DELETE FROM knowledge_vectors WHERE doc_id = ?", (doc_id,))
    db.commit()

    record_operation_log(username, '删除', '知识库文档', f'删除文档ID：{doc_id}')

    return jsonify({
        'code': 200,
        'message': '删除成功',
        'data': None
    })


@knowledge_ext_bp.route('/api/knowledge/documents/upload', methods=['POST'])
@token_required
def upload_documents():
    """批量上传文档文件。"""
    data = request.current_user
    username = data.get('username', '')

    if 'files' not in request.files:
        return jsonify({'code': 400, 'message': '请上传文件', 'data': None})

    files = request.files.getlist('files')
    doc_type = request.form.get('doc_type', 'other')
    cust_id = request.form.get('cust_id', type=int)
    business_id = request.form.get('business_id', type=int)
    tags = request.form.get('tags', '')

    batch_id = datetime.now().strftime('%Y%m%d%H%M%S') + '_' + uuid.uuid4().hex[:6]
    results = []
    db = get_db()
    cursor = db.cursor()

    for file in files:
        if not file or not file.filename:
            continue

        if not _allowed_file(file.filename):
            results.append({
                'filename': file.filename,
                'status': 'failed',
                'reason': '不支持的文件格式'
            })
            continue

        # secure_filename 对中文会返回空字符串，需要保留原始文件名
        original_filename = file.filename
        safe_filename = secure_filename(original_filename)
        if not safe_filename:
            # 中文文件名：用时间戳+保留扩展名
            ext = os.path.splitext(original_filename)[1] if '.' in original_filename else ''
            safe_filename = f"upload_{datetime.now().strftime('%Y%m%d%H%M%S')}_{uuid.uuid4().hex[:6]}{ext}"
        unique_name = f"{datetime.now().strftime('%Y%m%d%H%M%S')}_{safe_filename}"
        file_path = os.path.join(UPLOAD_DIR, unique_name)

        file.save(file_path)
        file_size = os.path.getsize(file_path)

        content = _parse_file(file_path, file.mimetype)
        if content is None:
            content = f"[文件 {original_filename} 无法解析，请安装相应依赖库]"

        title = os.path.splitext(original_filename)[0]

        cursor.execute("""
            INSERT INTO knowledge_documents
            (doc_type, title, content, file_path, file_name, file_size, mime_type,
             cust_id, business_id, import_batch_id, owner_id, tags, processed)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, 0)
        """, (doc_type, title, content, file_path, original_filename, file_size, file.mimetype,
              cust_id, business_id, batch_id, username, tags))
        db.commit()

        doc_id = cursor.lastrowid

        try:
            if content and not content.startswith('[文件'):
                index_document(doc_id, f"{title}\n{content}")
                cursor.execute("UPDATE knowledge_documents SET processed = 1, processed_at = ? WHERE id = ?",
                               (datetime.now().strftime('%Y-%m-%d %H:%M:%S'), doc_id))
                db.commit()
                results.append({'filename': original_filename, 'status': 'success', 'doc_id': doc_id})
            else:
                results.append({'filename': original_filename, 'status': 'parsed_with_notes', 'doc_id': doc_id})
        except Exception as e:
            results.append({'filename': original_filename, 'status': 'indexed_failed', 'doc_id': doc_id})

    record_operation_log(username, '批量导入', '知识库', f'批次{batch_id} 上传{len(results)}个文件')

    return jsonify({
        'code': 200,
        'message': f'上传完成，成功{sum(1 for r in results if r["status"] in ("success", "parsed_with_notes"))}个',
        'data': {
            'batch_id': batch_id,
            'total': len(files),
            'results': results
        }
    })


@knowledge_ext_bp.route('/api/knowledge/documents/batch-import', methods=['POST'])
@token_required
def batch_import():
    """批量文本导入（不通过文件上传，直接提供文本内容）。"""
    data = request.current_user
    username = data.get('username', '')
    req_data = request.get_json(silent=True) or {}

    documents = req_data.get('documents', [])
    if not documents:
        return jsonify({'code': 400, 'message': '请提供文档列表', 'data': None})

    batch_id = datetime.now().strftime('%Y%m%d%H%M%S') + '_' + uuid.uuid4().hex[:6]
    db = get_db()
    cursor = db.cursor()
    results = []

    for doc in documents:
        doc_type = doc.get('doc_type', 'other')
        title = doc.get('title', '').strip()
        content = doc.get('content', '').strip()
        cust_id = doc.get('cust_id')
        business_id = doc.get('business_id')
        tags = doc.get('tags', '')

        if not title:
            results.append({'title': title, 'status': 'skipped', 'reason': '标题为空'})
            continue

        cursor.execute("""
            INSERT INTO knowledge_documents
            (doc_type, title, content, cust_id, business_id, import_batch_id,
             owner_id, tags, processed)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, 1)
        """, (doc_type, title, content, cust_id, business_id, batch_id, username, tags))
        db.commit()

        doc_id = cursor.lastrowid

        try:
            if content:
                index_document(doc_id, f"{title}\n{content}")
        except Exception:
            pass

        results.append({'title': title, 'status': 'success', 'doc_id': doc_id})

    record_operation_log(username, '批量导入', '知识库', f'批次{batch_id} 导入{len(results)}条')

    return jsonify({
        'code': 200,
        'message': f'导入完成，成功{sum(1 for r in results if r["status"] == "success")}条',
        'data': {
            'batch_id': batch_id,
            'total': len(documents),
            'results': results
        }
    })


@knowledge_ext_bp.route('/api/knowledge/documents/search', methods=['POST'])
@token_required
def search_documents():
    """语义/混合搜索知识库文档。"""
    data = request.current_user
    req_data = request.get_json(silent=True) or {}

    query = req_data.get('query', '').strip()
    search_type = req_data.get('search_type', 'hybrid')
    doc_type = req_data.get('doc_type')
    cust_id = req_data.get('cust_id')
    business_id = req_data.get('business_id')
    top_k = req_data.get('top_k', 10)

    if not query:
        return jsonify({'code': 400, 'message': '请提供搜索词', 'data': None})

    if search_type == 'semantic':
        results = semantic_search(query, doc_type=doc_type, cust_id=cust_id,
                                  business_id=business_id, top_k=top_k)
    elif search_type == 'hybrid':
        filters = {}
        if doc_type:
            filters['doc_type'] = doc_type
        if cust_id:
            filters['cust_id'] = cust_id
        if business_id:
            filters['business_id'] = business_id
        results = hybrid_search(query, filters=filters if filters else None, top_k=top_k)
    else:
        results = semantic_search(query, doc_type=doc_type, top_k=top_k)

    for r in results:
        r['doc_type_display'] = DOC_TYPE_MAP.get(r.get('doc_type', ''), r.get('doc_type', ''))

    return jsonify({
        'code': 200,
        'message': 'success',
        'data': {
            'query': query,
            'search_type': search_type,
            'results': results,
            'total': len(results)
        }
    })


@knowledge_ext_bp.route('/api/knowledge/sync', methods=['POST'])
@token_required
def sync_crm_data():
    """CRM数据自动同步到知识库。"""
    data = request.current_user
    username = data.get('username', '')
    req_data = request.get_json(silent=True) or {}

    modules = req_data.get('modules', ['customers', 'business', 'contracts', 'visits'])
    db = get_db()
    cursor = db.cursor()
    sync_results = {}

    if 'customers' in modules:
        try:
            cursor.execute("""
                SELECT c.id, c.company, c.industry, c.level, c.region,
                       c.address, c.contact, c.phone, c.email, c.website,
                       c.notes, c.created_at
                FROM customers c
            """)
            customers = cursor.fetchall()
            synced = 0
            for cust in customers:
                cust_dict = dict(cust)
                cursor.execute("""
                    INSERT OR IGNORE INTO knowledge_documents
                    (doc_type, title, content, cust_id, owner_id, tags, processed)
                    VALUES (?, ?, ?, ?, ?, ?, 1)
                """, (
                    'customer_info',
                    f"客户资料：{cust_dict['company']}",
                    json.dumps(cust_dict, ensure_ascii=False),
                    cust_dict['id'],
                    'system',
                    f"客户,{cust_dict.get('industry', '')}"
                ))
                if cursor.rowcount > 0:
                    synced += 1
            db.commit()
            sync_results['customers'] = f'同步{synced}条客户资料'
        except Exception as e:
            sync_results['customers'] = f'同步失败：{str(e)}'

    if 'business' in modules:
        try:
            cursor.execute("""
                SELECT b.id, b.title, b.stage, b.probability, b.amount,
                       b.industry, b.cust_id, b.owner_id, b.description,
                       b.created_at
                FROM business b
                WHERE b.status = 'active'
            """)
            business_list = cursor.fetchall()
            synced = 0
            for biz in business_list:
                biz_dict = dict(biz)
                cursor.execute("""
                    INSERT OR IGNORE INTO knowledge_documents
                    (doc_type, title, content, cust_id, business_id, owner_id, tags, processed)
                    VALUES (?, ?, ?, ?, ?, ?, ?, 1)
                """, (
                    'other',
                    f"商机：{biz_dict['title']}",
                    json.dumps(biz_dict, ensure_ascii=False),
                    biz_dict.get('cust_id'),
                    biz_dict['id'],
                    biz_dict.get('owner_id', 'system'),
                    f"商机,{biz_dict.get('stage', '')},{biz_dict.get('industry', '')}"
                ))
                if cursor.rowcount > 0:
                    synced += 1
            db.commit()
            sync_results['business'] = f'同步{synced}条商机'
        except Exception as e:
            sync_results['business'] = f'同步失败：{str(e)}'

    if 'contracts' in modules:
        try:
            cursor.execute("""
                SELECT c.id, c.contract_no, c.title, c.amount, c.status,
                       c.cust_id, c.owner_id, c.sign_date, c.description
                FROM contracts c
                WHERE c.status IN ('执行中', '已完成')
            """)
            contracts = cursor.fetchall()
            synced = 0
            for contract in contracts:
                contract_dict = dict(contract)
                cursor.execute("""
                    INSERT OR IGNORE INTO knowledge_documents
                    (doc_type, title, content, cust_id, contract_id, owner_id, tags, processed)
                    VALUES (?, ?, ?, ?, ?, ?, ?, 1)
                """, (
                    'contract',
                    f"合同：{contract_dict['title'] or contract_dict['contract_no']}",
                    json.dumps(contract_dict, ensure_ascii=False),
                    contract_dict.get('cust_id'),
                    contract_dict['id'],
                    contract_dict.get('owner_id', 'system'),
                    f"合同,{contract_dict.get('status', '')}"
                ))
                if cursor.rowcount > 0:
                    synced += 1
            db.commit()
            sync_results['contracts'] = f'同步{synced}条合同'
        except Exception as e:
            sync_results['contracts'] = f'同步失败：{str(e)}'

    if 'visits' in modules:
        try:
            cursor.execute("""
                SELECT v.id, v.cust_id, v.visitor_id, v.plan_date, v.result,
                       v.purpose, v.notes, v.work_content
                FROM visits v
                WHERE v.actual_date IS NOT NULL
            """)
            visits = cursor.fetchall()
            synced = 0
            for visit in visits:
                visit_dict = dict(visit)
                content_parts = []
                for key in ['purpose', 'result', 'notes', 'work_content']:
                    if visit_dict.get(key):
                        content_parts.append(f"{key}: {visit_dict[key]}")
                content = '\n'.join(content_parts) if content_parts else ''

                cursor.execute("""
                    INSERT OR IGNORE INTO knowledge_documents
                    (doc_type, title, content, cust_id, owner_id, tags, processed)
                    VALUES (?, ?, ?, ?, ?, ?, 1)
                """, (
                    'visit_summary',
                    f"拜访纪要：{visit_dict.get('plan_date', '')} {visit_dict.get('purpose', '')}",
                    content or json.dumps(visit_dict, ensure_ascii=False),
                    visit_dict.get('cust_id'),
                    visit_dict.get('visitor_id', 'system'),
                    '拜访纪要'
                ))
                if cursor.rowcount > 0:
                    synced += 1
            db.commit()
            sync_results['visits'] = f'同步{synced}条拜访纪要'
        except Exception as e:
            sync_results['visits'] = f'同步失败：{str(e)}'

    try:
        cursor.execute("""
            INSERT INTO crm_sync_config (module, last_sync_at, sync_interval_hours)
            VALUES (?, ?, ?)
            ON CONFLICT(module) DO UPDATE SET last_sync_at = excluded.last_sync_at
        """, (
            ','.join(modules),
            datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
            24
        ))
        db.commit()
    except Exception:
        pass

    record_operation_log(username, '数据同步', '知识库', f'同步模块：{", ".join(modules)}')

    return jsonify({
        'code': 200,
        'message': '同步完成',
        'data': {
            'sync_results': sync_results,
            'synced_at': datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        }
    })


@knowledge_ext_bp.route('/api/knowledge/personnel-qualifications', methods=['GET'])
@token_required
def get_personnel_qualifications():
    """获取人员资质列表。"""
    data = request.current_user
    page = request.args.get('page', 1, type=int)
    per_page = request.args.get('per_page', 20, type=int)
    username = request.args.get('username', '')
    qualification_type = request.args.get('qualification_type', '')
    status = request.args.get('status', '')

    db = get_db()
    cursor = db.cursor()

    conditions = []
    params = []

    if username:
        conditions.append("username = ?")
        params.append(username)
    if qualification_type:
        conditions.append("qualification_type = ?")
        params.append(qualification_type)
    if status:
        conditions.append("status = ?")
        params.append(status)

    where_clause = "WHERE " + " AND ".join(conditions) if conditions else ""

    cursor.execute(f"SELECT COUNT(*) as total FROM personnel_qualifications {where_clause}", params)
    total = cursor.fetchone()['total']

    offset = (page - 1) * per_page
    cursor.execute(f"""
        SELECT * FROM personnel_qualifications
        {where_clause}
        ORDER BY created_at DESC
        LIMIT ? OFFSET ?
    """, params + [per_page, offset])

    quals = [dict(row) for row in cursor.fetchall()]

    return jsonify({
        'code': 200,
        'message': 'success',
        'data': {
            'items': quals,
            'total': total,
            'page': page,
            'per_page': per_page
        }
    })


@knowledge_ext_bp.route('/api/knowledge/personnel-qualifications', methods=['POST'])
@token_required
def create_personnel_qualification():
    """创建人员资质。"""
    data = request.current_user
    username = data.get('username', '')
    req_data = request.get_json(silent=True) or {}

    required = ['username', 'name', 'qualification_type']
    for field in required:
        if not req_data.get(field):
            return jsonify({'code': 400, 'message': f'{field} 为必填项', 'data': None})

    db = get_db()
    cursor = db.cursor()

    cursor.execute("""
        INSERT INTO personnel_qualifications
        (username, name, qualification_type, qualification_name, certificate_no,
         issue_date, expire_date, issue_authority, specialty, level, status)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, '有效')
    """, (
        req_data['username'], req_data['name'],
        req_data['qualification_type'],
        req_data.get('qualification_name', ''),
        req_data.get('certificate_no', ''),
        req_data.get('issue_date'),
        req_data.get('expire_date'),
        req_data.get('issue_authority', ''),
        req_data.get('specialty', ''),
        req_data.get('level', '')
    ))
    db.commit()

    qual_id = cursor.lastrowid

    record_operation_log(username, '创建', '人员资质', f'创建{req_data["name"]}的资质记录')

    return jsonify({
        'code': 200,
        'message': '创建成功',
        'data': {'id': qual_id}
    })


@knowledge_ext_bp.route('/api/knowledge/personnel-qualifications/<int:qual_id>', methods=['PUT'])
@token_required
def update_personnel_qualification(qual_id):
    """更新人员资质。"""
    data = request.current_user
    username = data.get('username', '')
    req_data = request.get_json(silent=True) or {}

    db = get_db()
    cursor = db.cursor()

    cursor.execute("SELECT id FROM personnel_qualifications WHERE id = ?", (qual_id,))
    if not cursor.fetchone():
        return jsonify({'code': 404, 'message': '资质记录不存在', 'data': None})

    fields = []
    params = []
    for field in ('qualification_type', 'qualification_name', 'certificate_no',
                  'issue_date', 'expire_date', 'issue_authority', 'specialty', 'level', 'status'):
        if field in req_data:
            fields.append(f"{field} = ?")
            params.append(req_data[field])

    if not fields:
        return jsonify({'code': 400, 'message': '无更新字段', 'data': None})

    fields.append("updated_at = CURRENT_TIMESTAMP")
    params.append(qual_id)

    cursor.execute(f"UPDATE personnel_qualifications SET {', '.join(fields)} WHERE id = ?", params)
    db.commit()

    record_operation_log(username, '更新', '人员资质', f'更新资质ID：{qual_id}')

    return jsonify({'code': 200, 'message': '更新成功', 'data': None})


@knowledge_ext_bp.route('/api/knowledge/personnel-qualifications/<int:qual_id>', methods=['DELETE'])
@token_required
def delete_personnel_qualification(qual_id):
    """删除人员资质。"""
    data = request.current_user
    username = data.get('username', '')

    db = get_db()
    cursor = db.cursor()

    cursor.execute("SELECT id FROM personnel_qualifications WHERE id = ?", (qual_id,))
    if not cursor.fetchone():
        return jsonify({'code': 404, 'message': '资质记录不存在', 'data': None})

    cursor.execute("DELETE FROM personnel_qualifications WHERE id = ?", (qual_id,))
    db.commit()

    record_operation_log(username, '删除', '人员资质', f'删除资质ID：{qual_id}')

    return jsonify({'code': 200, 'message': '删除成功', 'data': None})


@knowledge_ext_bp.route('/api/knowledge/company-qualifications', methods=['GET'])
@token_required
def get_company_qualifications():
    """获取企业资质列表。"""
    data = request.current_user
    page = request.args.get('page', 1, type=int)
    per_page = request.args.get('per_page', 20, type=int)
    qualification_type = request.args.get('qualification_type', '')
    status = request.args.get('status', '')

    db = get_db()
    cursor = db.cursor()

    conditions = []
    params = []

    if qualification_type:
        conditions.append("qualification_type = ?")
        params.append(qualification_type)
    if status:
        conditions.append("status = ?")
        params.append(status)

    where_clause = "WHERE " + " AND ".join(conditions) if conditions else ""

    cursor.execute(f"SELECT COUNT(*) as total FROM company_qualifications {where_clause}", params)
    total = cursor.fetchone()['total']

    offset = (page - 1) * per_page
    cursor.execute(f"""
        SELECT * FROM company_qualifications
        {where_clause}
        ORDER BY created_at DESC
        LIMIT ? OFFSET ?
    """, params + [per_page, offset])

    quals = [dict(row) for row in cursor.fetchall()]

    return jsonify({
        'code': 200,
        'message': 'success',
        'data': {
            'items': quals,
            'total': total,
            'page': page,
            'per_page': per_page
        }
    })


@knowledge_ext_bp.route('/api/knowledge/company-qualifications', methods=['POST'])
@token_required
def create_company_qualification():
    """创建企业资质。"""
    data = request.current_user
    username = data.get('username', '')
    req_data = request.get_json(silent=True) or {}

    if not req_data.get('qualification_type'):
        return jsonify({'code': 400, 'message': 'qualification_type 为必填项', 'data': None})

    db = get_db()
    cursor = db.cursor()

    cursor.execute("""
        INSERT INTO company_qualifications
        (qualification_type, qualification_name, certificate_no, issue_date,
         expire_date, issue_authority, scope, level, status)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, '有效')
    """, (
        req_data['qualification_type'],
        req_data.get('qualification_name', ''),
        req_data.get('certificate_no', ''),
        req_data.get('issue_date'),
        req_data.get('expire_date'),
        req_data.get('issue_authority', ''),
        req_data.get('scope', ''),
        req_data.get('level', '')
    ))
    db.commit()

    qual_id = cursor.lastrowid

    record_operation_log(username, '创建', '企业资质', f'创建{req_data.get("qualification_name") or req_data["qualification_type"]}')

    return jsonify({'code': 200, 'message': '创建成功', 'data': {'id': qual_id}})


@knowledge_ext_bp.route('/api/knowledge/company-qualifications/<int:qual_id>', methods=['PUT'])
@token_required
def update_company_qualification(qual_id):
    """更新企业资质。"""
    data = request.current_user
    username = data.get('username', '')
    req_data = request.get_json(silent=True) or {}

    db = get_db()
    cursor = db.cursor()

    cursor.execute("SELECT id FROM company_qualifications WHERE id = ?", (qual_id,))
    if not cursor.fetchone():
        return jsonify({'code': 404, 'message': '资质记录不存在', 'data': None})

    fields = []
    params = []
    for field in ('qualification_type', 'qualification_name', 'certificate_no',
                  'issue_date', 'expire_date', 'issue_authority', 'scope', 'level', 'status'):
        if field in req_data:
            fields.append(f"{field} = ?")
            params.append(req_data[field])

    if not fields:
        return jsonify({'code': 400, 'message': '无更新字段', 'data': None})

    fields.append("updated_at = CURRENT_TIMESTAMP")
    params.append(qual_id)

    cursor.execute(f"UPDATE company_qualifications SET {', '.join(fields)} WHERE id = ?", params)
    db.commit()

    record_operation_log(username, '更新', '企业资质', f'更新资质ID：{qual_id}')

    return jsonify({'code': 200, 'message': '更新成功', 'data': None})


@knowledge_ext_bp.route('/api/knowledge/company-qualifications/<int:qual_id>', methods=['DELETE'])
@token_required
def delete_company_qualification(qual_id):
    """删除企业资质。"""
    data = request.current_user
    username = data.get('username', '')

    db = get_db()
    cursor = db.cursor()

    cursor.execute("SELECT id FROM company_qualifications WHERE id = ?", (qual_id,))
    if not cursor.fetchone():
        return jsonify({'code': 404, 'message': '资质记录不存在', 'data': None})

    cursor.execute("DELETE FROM company_qualifications WHERE id = ?", (qual_id,))
    db.commit()

    record_operation_log(username, '删除', '企业资质', f'删除资质ID：{qual_id}')

    return jsonify({'code': 200, 'message': '删除成功', 'data': None})


@knowledge_ext_bp.route('/api/knowledge/rebuild-vectors', methods=['POST'])
@token_required
def rebuild_vectors():
    """重建所有向量索引。"""
    data = request.current_user
    username = data.get('username', '')

    try:
        stats = rebuild_all_vectors()
        record_operation_log(username, '重建', '向量索引', f'重建完成，处理{stats.get("processed", 0)}条')
        return jsonify({
            'code': 200,
            'message': '重建完成',
            'data': stats
        })
    except Exception as e:
        return jsonify({'code': 500, 'message': f'重建失败：{str(e)}', 'data': None})


@knowledge_ext_bp.route('/api/knowledge/stats', methods=['GET'])
@token_required
def get_knowledge_stats():
    """获取知识库统计信息。"""
    db = get_db()
    cursor = db.cursor()

    cursor.execute("SELECT doc_type, COUNT(*) as count FROM knowledge_documents GROUP BY doc_type")
    doc_type_counts = [dict(row) for row in cursor.fetchall()]

    cursor.execute("SELECT COUNT(*) as total FROM knowledge_documents WHERE processed = 1")
    processed_count = cursor.fetchone()['total']

    cursor.execute("SELECT COUNT(*) as total FROM knowledge_documents")
    total_docs = cursor.fetchone()['total']

    cursor.execute("SELECT COUNT(*) as total FROM knowledge_vectors")
    total_vectors = cursor.fetchone()['total']

    cursor.execute("SELECT COUNT(*) as total FROM personnel_qualifications WHERE status = '有效'")
    personnel_count = cursor.fetchone()['total']

    cursor.execute("SELECT COUNT(*) as total FROM company_qualifications WHERE status = '有效'")
    company_count = cursor.fetchone()['total']

    cursor.execute("SELECT COUNT(*) as total FROM knowledge_documents WHERE analysis_status = 'completed'")
    analyzed_count = cursor.fetchone()['total']

    return jsonify({
        'code': 200,
        'message': 'success',
        'data': {
            'total_documents': total_docs,
            'processed_documents': processed_count,
            'total_vectors': total_vectors,
            'analyzed_documents': analyzed_count,
            'personnel_qualifications': personnel_count,
            'company_qualifications': company_count,
            'doc_type_distribution': [
                {'type': item['doc_type'], 'count': item['count'], 'label': DOC_TYPE_MAP.get(item['doc_type'], item['doc_type'])}
                for item in doc_type_counts
            ]
        }
    })


@knowledge_ext_bp.route('/api/knowledge/documents/<int:doc_id>/analyze', methods=['POST'])
@token_required
def analyze_single_document(doc_id):
    """分析单个文档，返回结构化分析结果。
    三步式：短连接读文档 → LLM分析(不持锁) → 短事务写结果
    """
    data = request.current_user
    username = data.get('username', '')

    # 第1步：短事务读取文档内容（SELECT 不持有写锁，读取后连接可复用）
    db = get_db()
    cursor = db.cursor()
    cursor.execute("SELECT title, content, doc_type FROM knowledge_documents WHERE id = ?", (doc_id,))
    doc = cursor.fetchone()
    if not doc:
        return jsonify({'code': 404, 'message': '文档不存在', 'data': None})

    doc_dict = dict(doc)
    # 第2步：LLM分析（SELECT 未持有写锁，不影响其他请求的写入）
    analysis = analyze_document(
        doc_dict.get('title', ''),
        doc_dict.get('content', ''),
        doc_dict.get('doc_type', 'other')
    )

    # 第3步：短事务写结果（毫秒级写锁）
    if analysis:
        analysis_json = json.dumps(analysis, ensure_ascii=False)
        cursor.execute("""
            UPDATE knowledge_documents
            SET analysis_result = ?, analysis_status = 'completed',
                analyzed_at = ?, updated_at = ?
            WHERE id = ?
        """, (analysis_json, analysis.get('analyzed_at', datetime.now().strftime('%Y-%m-%d %H:%M:%S')),
              datetime.now().strftime('%Y-%m-%d %H:%M:%S'), doc_id))
        db.commit()

        record_operation_log(username, '分析', '知识库文档', f'分析文档ID：{doc_id}')

        return jsonify({
            'code': 200,
            'message': '分析完成',
            'data': analysis
        })
    else:
        cursor.execute("""
            UPDATE knowledge_documents
            SET analysis_status = 'failed', updated_at = ?
            WHERE id = ?
        """, (datetime.now().strftime('%Y-%m-%d %H:%M:%S'), doc_id))
        db.commit()
        return jsonify({'code': 400, 'message': '文档内容为空，无法分析', 'data': None})


@knowledge_ext_bp.route('/api/knowledge/documents/batch-analyze', methods=['POST'])
@token_required
def batch_analyze_documents():
    """批量分析文档。
    三步式：短连接读取文档列表 → 逐个LLM分析(不持锁) → 每个分析后短事务写入
    """
    data = request.current_user
    username = data.get('username', '')
    req_data = request.get_json(silent=True) or {}

    doc_ids = req_data.get('doc_ids', [])
    # 第1步：短事务读取待分析文档列表
    db = get_db()
    cursor = db.cursor()
    if not doc_ids:
        cursor.execute("""
            SELECT id, title, content, doc_type FROM knowledge_documents
            WHERE (analysis_status IS NULL OR analysis_status != 'completed')
            AND content IS NOT NULL AND content != ''
            LIMIT 50
        """)
        docs_to_analyze = [dict(row) for row in cursor.fetchall()]
    else:
        placeholders = ','.join('?' * len(doc_ids[:50]))
        cursor.execute(f"SELECT id, title, content, doc_type FROM knowledge_documents WHERE id IN ({placeholders})",
                       doc_ids[:50])
        docs_to_analyze = [dict(row) for row in cursor.fetchall()]
        docs_to_analyze = [d for d in docs_to_analyze if d.get('content')]

    if not docs_to_analyze:
        return jsonify({'code': 200, 'message': '没有需要分析的文档', 'data': {'analyzed': 0}})

    # 第2步：逐个LLM分析（SELECT未持写锁，不影响其他请求；每个分析后立即commit短事务写入）
    results = []
    analyzed_count = 0
    for doc_dict in docs_to_analyze:
        doc_id = doc_dict['id']
        analysis = analyze_document(
            doc_dict.get('title', ''),
            doc_dict.get('content', ''),
            doc_dict.get('doc_type', 'other')
        )

        # 第3步：每个文档分析完后，短事务写入并commit（毫秒级写锁）
        if analysis:
            analysis_json = json.dumps(analysis, ensure_ascii=False)
            cursor.execute("""
                UPDATE knowledge_documents
                SET analysis_result = ?, analysis_status = 'completed',
                    analyzed_at = ?, updated_at = ?
                WHERE id = ?
            """, (analysis_json, analysis.get('analyzed_at', datetime.now().strftime('%Y-%m-%d %H:%M:%S')),
                  datetime.now().strftime('%Y-%m-%d %H:%M:%S'), doc_id))
            db.commit()
            analyzed_count += 1
            results.append({'id': doc_id, 'title': doc_dict.get('title', ''), 'analysis': analysis})

    record_operation_log(username, '批量分析', '知识库文档', f'分析{analyzed_count}个文档')

    return jsonify({
        'code': 200,
        'message': f'分析完成，共分析 {analyzed_count} 个文档',
        'data': {
            'analyzed': analyzed_count,
            'results': results
        }
    })


@knowledge_ext_bp.route('/api/knowledge/documents/<int:doc_id>/analysis', methods=['GET'])
@token_required
def get_document_analysis(doc_id):
    """获取文档的分析结果。"""
    db = get_db()
    cursor = db.cursor()

    cursor.execute("SELECT analysis_result, analysis_status, analyzed_at FROM knowledge_documents WHERE id = ?", (doc_id,))
    doc = cursor.fetchone()
    if not doc:
        return jsonify({'code': 404, 'message': '文档不存在', 'data': None})

    analysis = None
    if doc['analysis_result']:
        try:
            analysis = json.loads(doc['analysis_result'])
        except:
            analysis = None

    return jsonify({
        'code': 200,
        'message': 'success',
        'data': {
            'analysis': analysis,
            'analysis_status': doc['analysis_status'] or 'pending',
            'analyzed_at': doc['analyzed_at']
        }
    })


def register_routes(app):
    app.register_blueprint(knowledge_ext_bp)