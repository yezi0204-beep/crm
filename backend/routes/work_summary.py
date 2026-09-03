"""应用中心工作总结模块

基于部门（默认应用中心）在指定周期内的真实业务数据，多维度聚合：
  1. 客户拜访与关系维护（visits）
  2. 商机推进与项目跟进（business）
  3. 签约与回款（contracts / payment_records）
  4. 客户覆盖（customers）
通过 LLM 生成结构化总结报告（Markdown），并支持导出排版美观的 Word（python-docx）。

路由：
  GET  /api/work-summary/data       周期数据聚合预览
  POST /api/work-summary/generate   LLM 生成总结（长请求）
  POST /api/work-summary/export-word 导出 Word（长请求，返回文件流）
"""
import json
import io
import re
from datetime import datetime, timedelta, date

from flask import request, jsonify, send_file
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls, qn
from docx.oxml import parse_xml

from extensions import get_db, token_required, record_operation_log
from qa_engine import call_llm

from . import work_summary_bp

DEFAULT_DEPARTMENT = '应用中心'
THEME_DARK = RGBColor(0x1F, 0x4E, 0x79)    # 深蓝（标题）
THEME_MID = RGBColor(0x2E, 0x75, 0xB6)     # 中蓝（三级标题）
THEME_GRAY = RGBColor(0x59, 0x59, 0x59)    # 灰色（副标题/说明）


# ==================== 周期与数据聚合 ====================

def _period_range(period_type, offset):
    """计算周期起止日期（含首尾）。offset=0 当前周期，1 上一个周期。"""
    today = date.today()
    if period_type == 'month':
        offset_month = today.month - int(offset or 0)
        year, month = today.year, offset_month
        while month <= 0:
            year -= 1
            month += 12
        start = date(year, month, 1)
        end = date(year + (month // 12), (month % 12) + 1, 1) - timedelta(days=1)
    else:  # week：周一 ~ 周日
        this_monday = today - timedelta(days=today.weekday())
        start = this_monday - timedelta(days=7 * int(offset or 0))
        end = start + timedelta(days=6)
    return start.strftime('%Y-%m-%d'), end.strftime('%Y-%m-%d')


def _period_label(period_type, offset):
    start, end = _period_range(period_type, offset)
    if period_type == 'month':
        return start[:7].replace('-', '年') + '月'
    return f'{start} ~ {end}（第{date(*map(int, start.split("-"))).isocalendar()[1]}周）'


def _collect_data(cursor, department, start, end):
    """聚合部门在周期内的多维度业务数据。SQL 全参数化。"""
    data = {'department': department, 'start': start, 'end': end}

    # 部门在职人员
    cursor.execute("""
        SELECT username, name, role FROM users
        WHERE department = ? AND (status = '在职' OR status IS NULL OR status = '')
        ORDER BY CASE role WHEN '主任' THEN 0 WHEN '院长' THEN 1 ELSE 2 END, name
    """, (department,))
    users = [dict(r) for r in cursor.fetchall()]
    data['users'] = [{'username': u['username'], 'name': u['name'], 'role': u['role']} for u in users]
    usernames = [u['username'] for u in users]
    ph = ','.join('?' * len(usernames)) if usernames else "''"

    # ===== 维度1：客户拜访（visits）=====
    cursor.execute(f"""
        SELECT v.visitor_id, u.name as visitor_name, v.plan_date, v.status,
               v.work_type, v.purpose, v.result, v.location, v.notes, v.work_content,
               c.company as customer_company, c.name as customer_name
        FROM visits v
        LEFT JOIN users u ON v.visitor_id = u.username
        LEFT JOIN customers c ON v.cust_id = c.id
        WHERE v.visitor_id IN ({ph}) AND v.plan_date >= ? AND v.plan_date <= ?
          AND v.status != 'cancelled'
        ORDER BY v.visitor_id, v.plan_date
    """, (*usernames, start, end))
    visits = [dict(r) for r in cursor.fetchall()]
    data['visits_total'] = len(visits)
    data['visits_completed'] = sum(1 for v in visits if v['status'] == 'completed')
    data['visits_planned'] = sum(1 for v in visits if v['status'] == 'planned')
    data['customers_covered'] = len({v['customer_company'] or v['customer_name'] for v in visits
                                     if v['customer_company'] or v['customer_name']})
    # 每人统计与明细（明细仅保留有内容的条目，控制 LLM 输入长度）
    per_user = []
    for u in users:
        uv = [v for v in visits if v['visitor_id'] == u['username']]
        details = [{
            'date': v['plan_date'], 'customer': v['customer_company'] or v['customer_name'] or '',
            'purpose': (v['purpose'] or v['work_content'] or '')[:80],
            'result': (v['result'] or '')[:80]
        } for v in uv]
        per_user.append({
            'name': u['name'], 'role': u['role'], 'total': len(uv),
            'completed': sum(1 for v in uv if v['status'] == 'completed'),
            'planned': sum(1 for v in uv if v['status'] == 'planned'),
            'details': details
        })
    data['visits_by_user'] = per_user

    # ===== 维度2：商机推进（business，以最近跟进时间 last_follow 衡量推进活动）=====
    cursor.execute(f"""
        SELECT owner_id, stage, COUNT(*) as cnt, COALESCE(SUM(amount), 0) as amount
        FROM business
        WHERE owner_id IN ({ph}) AND status NOT IN ('converted', 'rejected')
          AND date(COALESCE(NULLIF(last_follow, ''), created_at)) >= ?
          AND date(COALESCE(NULLIF(last_follow, ''), created_at)) <= ?
        GROUP BY owner_id, stage
    """, (*usernames, start, end))
    data['business_by_stage'] = [dict(r) for r in cursor.fetchall()]

    cursor.execute(f"""
        SELECT COUNT(*) as cnt, COALESCE(SUM(amount), 0) as amount
        FROM business
        WHERE owner_id IN ({ph}) AND date(created_at) >= ? AND date(created_at) <= ?
    """, (*usernames, start, end))
    row = cursor.fetchone()
    data['business_new'] = {'count': row['cnt'], 'amount_wan': round((row['amount'] or 0) / 10000, 2)}

    # ===== 维度3：签约与回款 =====
    cursor.execute(f"""
        SELECT COUNT(*) as cnt, COALESCE(SUM(total_amt), 0) as total
        FROM contracts
        WHERE owner_id IN ({ph}) AND date(sign_date) >= ? AND date(sign_date) <= ?
    """, (*usernames, start, end))
    row = cursor.fetchone()
    data['contracts_signed'] = {'count': row['cnt'], 'amount_yuan': round(row['total'] or 0, 2)}

    cursor.execute(f"""
        SELECT COUNT(*) as cnt, COALESCE(SUM(p.amount), 0) as total
        FROM payment_records p
        JOIN contracts ct ON p.contract_id = ct.id
        WHERE ct.owner_id IN ({ph}) AND date(p.payment_date) >= ? AND date(p.payment_date) <= ?
    """, (*usernames, start, end))
    row = cursor.fetchone()
    data['payments_received'] = {'count': row['cnt'], 'amount_yuan': round(row['total'] or 0, 2)}

    # ===== 维度4：客户覆盖（新增建档客户）=====
    cursor.execute(f"""
        SELECT COUNT(*) as cnt FROM customers
        WHERE owner_id IN ({ph}) AND date(created_at) >= ? AND date(created_at) <= ?
    """, (*usernames, start, end))
    data['customers_new'] = cursor.fetchone()['cnt']

    return data


# ==================== 接口：数据聚合预览 ====================

@work_summary_bp.route('/api/work-summary/data', methods=['GET'])
@token_required
def summary_data():
    payload = request.current_user
    period_type = request.args.get('period_type', 'week')
    offset = int(request.args.get('offset', 0) or 0)
    department = request.args.get('department', DEFAULT_DEPARTMENT)
    if period_type not in ('week', 'month'):
        return jsonify({'code': 400, 'message': 'period_type 必须为 week 或 month', 'data': None})

    start, end = _period_range(period_type, offset)
    db = get_db()
    cursor = db.cursor()
    data = _collect_data(cursor, department, start, end)
    data['period_label'] = _period_label(period_type, offset)
    return jsonify({'code': 200, 'message': 'success', 'data': data})


# ==================== 接口：LLM 生成总结 ====================

@work_summary_bp.route('/api/work-summary/generate', methods=['POST'])
@token_required
def generate_summary():
    payload = request.current_user
    body = request.get_json(silent=True) or {}
    period_type = body.get('period_type', 'week')
    offset = int(body.get('offset', 0) or 0)
    department = body.get('department', DEFAULT_DEPARTMENT)
    extra_note = (body.get('extra_note') or '').strip()

    if period_type not in ('week', 'month'):
        return jsonify({'code': 400, 'message': 'period_type 必须为 week 或 month', 'data': None})

    start, end = _period_range(period_type, offset)
    db = get_db()
    cursor = db.cursor()
    data = _collect_data(cursor, department, start, end)

    period_cn = '周' if period_type == 'week' else '月'
    period_label = _period_label(period_type, offset)

    # 关键指标摘要：置于最前，确保 LLM 一定引用这些数字（尤其是回款/签约）
    key_metrics = {
        '客户拜访总数': data['visits_total'],
        '已完成拜访': data['visits_completed'],
        '进行中拜访': data['visits_planned'],
        '覆盖客户数': data['customers_covered'],
        '新增商机数': data['business_new']['count'],
        '新增商机金额(万)': data['business_new']['amount_wan'],
        '签约合同数': data['contracts_signed']['count'],
        '签约合同金额(元)': data['contracts_signed']['amount_yuan'],
        '回款笔数': data['payments_received']['count'],
        '回款金额(元)': data['payments_received']['amount_yuan'],
        '新增建档客户数': data['customers_new'],
    }

    system_prompt = (
        '你是一名事业单位"应用中心"的工作总结撰写助手。请基于用户提供的真实业务数据撰写一份'
        f'部门{period_cn}工作总结，要求：\n'
        '1. 严禁编造任何数据、客户名或事件；数据为空的维度写"本周期暂无相关记录"，不得虚构。\n'
        '2. 【重要】输入最前方的"关键指标"是本周期核心数字，撰写"工作概览"与"签约与回款"两节时'
        '必须逐项准确引用这些数字（尤其是回款笔数与回款金额），不得写 0 或遗漏，除非指标本身为 0。\n'
        '3. 语言正式凝练，多用具体数字与客户/项目名称支撑，避免空话套话。\n'
        '4. 输出为 Markdown，严格采用以下结构（## 为二级标题）：\n'
        '## 一、工作概览\n'
        '## 二、主要工作与成果\n'
        '### （一）客户拜访与关系维护\n'
        '### （二）商机推进与项目跟进\n'
        '### （三）签约与回款\n'
        '## 三、数据统计\n'
        '（此节用 Markdown 表格展示每人拜访完成情况：姓名/角色/拜访总数/已完成/进行中）\n'
        '## 四、存在问题与改进方向\n'
        '## 五、下阶段工作计划\n'
        '5. 不要输出一级标题，从二级标题开始；不要使用代码块包裹全文。'
    )
    user_prompt = {
        '报告类型': f'{department}{period_cn}工作总结',
        '统计周期': f'{start} 至 {end}',
        '关键指标': key_metrics,
        '业务数据': data,
        '用户补充说明': extra_note or '无'
    }

    messages = [
        {'role': 'system', 'content': system_prompt},
        {'role': 'user', 'content': json.dumps(user_prompt, ensure_ascii=False, default=str)}
    ]
    answer = call_llm(messages, max_tokens=18000, timeout=360, enable_thinking=False)
    if not answer:
        return jsonify({'code': 500, 'message': 'LLM 服务未启用或调用失败，无法生成总结', 'data': None})

    record_operation_log(payload['username'], '生成', '工作总结',
                         f'{department} {period_label} 工作总结（LLM）')
    return jsonify({
        'code': 200, 'message': 'success',
        'data': {'content': answer, 'period_label': period_label,
                 'department': department, 'start': start, 'end': end}
    })


# ==================== Word 导出 ====================

_MD_TABLE_RE = re.compile(r'^\s*\|(.+)\|\s*$')
_MD_BOLD_RE = re.compile(r'\*\*(.+?)\*\*')


def _set_cn_font(run, name='宋体', size=12, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    # 中文字体需设置 eastAsia
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)


def _add_bottom_border(paragraph, color='8EAADB', size='12'):
    pPr = paragraph._p.get_or_add_pPr()
    pPr.append(parse_xml(
        f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="{size}" '
        f'w:space="4" w:color="{color}"/></w:pBdr>'))


def _shade_cell(cell, hex_color):
    cell._tc.get_or_add_tcPr().append(
        parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="{hex_color}"/>'))


def _fill_cell(cell, text, bold=False, white=False, size=10.5):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(str(text))
    _set_cn_font(run, name='微软雅黑', size=size, bold=bold,
                 color=RGBColor(0xFF, 0xFF, 0xFF) if white else None)


def _add_markdown_runs(paragraph, text, size, name='宋体', color=None, bold_all=False):
    """解析行内 **粗体**，向段落追加 run。"""
    pos = 0
    for m in _MD_BOLD_RE.finditer(text):
        if m.start() > pos:
            _set_cn_font(paragraph.add_run(text[pos:m.start()]), name=name, size=size, color=color, bold=bold_all)
        _set_cn_font(paragraph.add_run(m.group(1)), name=name, size=size, color=color, bold=True)
        pos = m.end()
    if pos < len(text):
        _set_cn_font(paragraph.add_run(text[pos:]), name=name, size=size, color=color, bold=bold_all)


def _add_stats_overview(doc, stats):
    """数据概览卡片表：每行 4 格（指标/值/指标/值），表头深蓝。"""
    items = [(s['label'], s['value']) for s in (stats or []) if s.get('label')]
    if not items:
        return
    rows = (len(items) + 1) // 2
    table = doc.add_table(rows=rows, cols=4)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i in range(rows):
        for j in range(2):
            idx = i * 2 + j
            if idx < len(items):
                label, value = items[idx]
                _fill_cell(table.cell(i, j * 2), label, bold=True)
                _shade_cell(table.cell(i, j * 2), 'DEEBF7')
                _fill_cell(table.cell(i, j * 2 + 1), value)
            else:
                _fill_cell(table.cell(i, j * 2), '')
                _fill_cell(table.cell(i, j * 2 + 1), '')
    doc.add_paragraph()


def _add_md_table(doc, header_cells, data_rows):
    table = doc.add_table(rows=len(data_rows) + 1, cols=len(header_cells))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(header_cells):
        _fill_cell(table.cell(0, j), h.strip(), bold=True, white=True)
        _shade_cell(table.cell(0, j), '4472C4')
    for i, row in enumerate(data_rows, 1):
        for j in range(len(header_cells)):
            v = row[j].strip() if j < len(row) else ''
            _fill_cell(table.cell(i, j), v if v else '-')
            if i % 2 == 0:
                _shade_cell(table.cell(i, j), 'F2F7FC')
    doc.add_paragraph()


def build_summary_docx(title, subtitle, stats, content_md):
    """构建美观的工作总结 Word 文档。"""
    doc = Document()

    # 页面设置：A4，适中边距
    section = doc.sections[0]
    section.page_width, section.page_height = Cm(21.0), Cm(29.7)
    section.top_margin = section.bottom_margin = Cm(2.3)
    section.left_margin = section.right_margin = Cm(2.6)

    # 默认正文字体
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(12)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 页脚页码（域代码 PAGE）
    footer_p = section.footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fld = parse_xml(f'<w:fldSimple {nsdecls("w")} w:instr="PAGE"/>')
    run_el = parse_xml(f'<w:r {nsdecls("w")}><w:rPr><w:rFonts w:eastAsia="宋体"/><w:sz w:val="18"/></w:rPr></w:r>')
    fld.append(run_el)
    footer_p._p.append(fld)

    # ===== 标题区 =====
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    _set_cn_font(p.add_run(title), name='微软雅黑', size=22, bold=True, color=THEME_DARK)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    _set_cn_font(p.add_run(subtitle), name='微软雅黑', size=10.5, color=THEME_GRAY)
    _add_bottom_border(p, color='4472C4', size='18')

    # ===== 数据概览 =====
    _add_stats_overview(doc, stats)
    doc.add_paragraph()

    # ===== 正文 Markdown 渲染 =====
    lines = (content_md or '').replace('\r\n', '\n').split('\n')
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        # markdown 表格（含分隔行）
        m = _MD_TABLE_RE.match(line)
        if m and i + 1 < len(lines) and re.match(r'^\s*\|[\s:\-|]+\|\s*$', lines[i + 1]):
            header = [c.strip() for c in m.group(1).split('|')]
            rows = []
            j = i + 2
            while j < len(lines) and _MD_TABLE_RE.match(lines[j]):
                rows.append(_MD_TABLE_RE.match(lines[j]).group(1).split('|'))
                j += 1
            _add_md_table(doc, header, rows)
            i = j
            continue

        if not line.strip():
            i += 1
            continue

        if line.startswith('### '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(4)
            _add_markdown_runs(p, line[4:], 13, name='微软雅黑', color=THEME_MID, bold_all=True)
        elif line.startswith('## '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(6)
            _add_markdown_runs(p, line[3:], 15, name='微软雅黑', color=THEME_DARK, bold_all=True)
            _add_bottom_border(p, color='BDD6EE', size='8')
        elif line.startswith('# '):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(10)
            _add_markdown_runs(p, line[2:], 16, name='微软雅黑', color=THEME_DARK, bold_all=True)
        elif re.match(r'^\s*[-*]\s+', line):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.74)
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            _set_cn_font(p.add_run('• '), size=12, color=THEME_MID, bold=True)
            _add_markdown_runs(p, re.sub(r'^\s*[-*]\s+', '', line), 12)
        elif re.match(r'^\s*\d+[.、)]\s+', line):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.74)
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            _add_markdown_runs(p, line.strip(), 12)
        else:
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Pt(24)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            _add_markdown_runs(p, line.strip(), 12)
        i += 1

    return doc


@work_summary_bp.route('/api/work-summary/export-word', methods=['POST'])
@token_required
def export_word():
    payload = request.current_user
    body = request.get_json(silent=True) or {}
    title = (body.get('title') or '工作总结').strip()
    subtitle = (body.get('subtitle') or '').strip()
    content = body.get('content') or ''
    stats = body.get('stats') or []

    if not content.strip():
        return jsonify({'code': 400, 'message': '总结内容为空，无法导出', 'data': None})

    doc = build_summary_docx(title, subtitle, stats, content)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    record_operation_log(payload['username'], '导出', '工作总结', f'导出 Word：{title}')
    safe_name = re.sub(r'[\\/:*?"<>|]', '_', title)
    return send_file(
        buf, as_attachment=True, download_name=f'{safe_name}.docx',
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')


def register_routes(app):
    app.register_blueprint(work_summary_bp)
